# -*- coding: utf-8 -*-
"""
@Software: PyCharm
@Author :张鹏
@Email  :1050035126@qq.com
@Date   :2020/07/08/0008 16:51
@Version:1.0
@Desc   :
解析word mysql表的结构定义
 输出mysql的插创建数据表的sql语句

"""
import docx
from docx import Document


def getTableNameList(doc):
    resultList = []

    paragraphs = doc.paragraphs
    for para in paragraphs:
        paraText = para.text
        if "表名" in paraText:
            resultList.append(paraText[3:].replace(" ", ""))

    return resultList


def getTableContentList(doc, tableNameList):
    result = {}

    tables = doc.tables
    tablesSize = len(tables)
    tableNameSize = len(tableNameList)
    # 这是数据表名称的说明表
    table0 = tables[0]
    tableNameExplainDic = {}
    for index, row in enumerate(table0.rows):
        if index == 0:
            continue

        cells = row.cells
        tableName = cells[1].text.replace(" ", "")
        tableExplain = cells[2].text.replace(" ", "")

        tableNameExplainDic[tableName] = tableExplain

    for tableIndex in range(tablesSize):
        if tableIndex == 0:
            continue
        rows = tables[tableIndex].rows
        rowSize = len(rows)

        tableName = tableNameList[tableIndex - 1]
        tempColList = []
        for rowIndex in range(rowSize):
            if rowIndex == 0:
                continue
            cells = rows[rowIndex].cells

            colIndex = cells[0].text
            name = cells[1].text
            comment = cells[2].text.replace(" ", "")
            dataType = cells[3].text
            dataLen = cells[4].text
            allowNullStr = cells[5].text

            if allowNullStr in ["","√"]:
                allowNull = ""
            elif allowNullStr == "自动递增":
                allowNull = "自动递增"
            else:
                allowNull = "NOT NULL"

            tempColList.append([name, comment, dataType, dataLen, allowNull])

        result[tableName] = tempColList

    return result, tableNameExplainDic


def outCreateTableSql(tableFormatDic, tableNameExplainDic):
    tableNameList = tableFormatDic.keys()

    for tableName in tableNameList:
        # 数据表的注释
        tableNameExplain = tableNameExplainDic.get(tableName)
        if not tableNameExplain:
            tableNameExplain = ""

        colList = tableFormatDic.get(tableName)
        print("\n")
        print(f"CREATE TABLE `{tableName}` (")
        # [name, comment, dataType, dataLen, allowNull]
        primaryKeySet = ""

        colSize = len(colList)
        tempSqlItemList = []
        for colIndex in range(colSize):
            col = colList[colIndex]

            name = col[0]
            comment = col[1]
            dataType = col[2]
            dataLen = col[3]
            allowNull = col[4]
            defaultNull = "DEFAULT NULL"

            # 简单校正错误
            if dataType == "datetime":
                dataLen = ""

            # 是否允许为空
            if allowNull:
                if allowNull == "自动递增":
                    allowNull = " NOT NULL AUTO_INCREMENT "
                else:
                    allowNull = f"{allowNull}"
            else:
                allowNull = ""

            # 添加数据长度定义
            if dataLen != "":
                dataLen = int(dataLen)
                dataLen = f"({dataLen})"

            if "主键" in comment:
                primaryKeySet = name

            # 对于主键和不允许空的 不设定默认值null
            if "主键" in comment or allowNull != "":
                defaultNull = ""

            tempSqlItemList.append(f"`{name}` {dataType}{dataLen} {allowNull} COMMENT '{comment}' {defaultNull} ")

        if primaryKeySet != "":
            tempSqlItemList.append(f"PRIMARY KEY (`{primaryKeySet}`)")

        print(",\n".join(tempSqlItemList))
        print(f") COMMENT='{tableNameExplain}' ENGINE=InnoDB DEFAULT CHARSET=utf8;")


if __name__ == '__main__':
    docPath = r"wordMysql.docx"

    doc = Document(docPath)

    tableNameList = getTableNameList(doc)
    tableFormatDic, tableNameExplainDic = getTableContentList(doc, tableNameList)
    outCreateTableSql(tableFormatDic, tableNameExplainDic)
